from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "templates" / "Z15_FPGA项目验收与测试报告.docx"
DST = ROOT / "reports" / "Z15_FPGA项目验收与测试报告_20260811_DMA最终版.docx"
from docx_table_geometry import (
    apply_table_geometry,
    column_widths_from_weights,
    section_content_width_dxa,
)


GREEN = RGBColor(0x00, 0x78, 0x3E)
BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"


def replace_paragraph(paragraph, text):
    if paragraph.runs:
        first = paragraph.runs[0]
        first.text = text
        for run in list(paragraph.runs[1:]):
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def set_cell(cell, text, *, bold=False, color=None, align=None, size=9):
    paragraph = cell.paragraphs[0]
    replace_paragraph(paragraph, text)
    run = paragraph.runs[0]
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "等线"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "等线")
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        paragraph.alignment = align
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_row(table, row_index, values):
    for cell, value in zip(table.rows[row_index].cells, values):
        set_cell(cell, value)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def style_new_table(table, widths, rows):
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        cells = table.rows[row_index].cells
        for col_index, value in enumerate(values):
            set_cell(
                cells[col_index], value,
                bold=(row_index == 0),
                color=(RGBColor(0xFF, 0xFF, 0xFF) if row_index == 0 else None),
                align=(WD_ALIGN_PARAGRAPH.CENTER if col_index > 0 or row_index == 0
                       else WD_ALIGN_PARAGRAPH.LEFT),
            )
            if row_index == 0:
                shade_cell(cells[col_index], BLUE)
            elif row_index % 2 == 0:
                shade_cell(cells[col_index], LIGHT_BLUE)
        if row_index > 0 and cells[-1].text == "PASS":
            cells[-1].paragraphs[0].runs[0].font.color.rgb = GREEN
            cells[-1].paragraphs[0].runs[0].bold = True
    repeat_header(table.rows[0])
    apply_table_geometry(table, widths)


doc = Document(SRC)

# Opening metadata and executive conclusion.
replace_paragraph(
    doc.paragraphs[1],
    "器件：XC7Z015CLG485-2｜最终 DMA/突发传输/批处理验收：2026-08-11",
)
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("FPGA 核心加速器已经形成"):
        replace_paragraph(
            paragraph,
            "FPGA 分子计算加速器已形成可烧录、可运行、可复测的完整闭环。"
            "在原 AXI-Lite 控制面和四类计算核心基础上，已加入 AXI DMA、HP0 突发传输、"
            "128-bit AXIS、64 项混合任务队列和批处理协议。最终实现时序、DRC、双向带宽、"
            "端到端性能、异常恢复、1000 批压力测试及旧接口兼容性全部通过。真实 RDKit/"
            "PyTorch/ADMET 模型一致性、TCP 后端和大规模业务数据仍依赖外部模型与系统输入，"
            "不属于本轮十二项 FPGA DMA 改造的完成条件。",
        )
        break

# Executive summary table.
t = doc.tables[0]
set_row(t, 2, ["100/150 MHz 时序与正式产物", "通过", "WNS +0.012/+0.460 ns；0 DRC/方法学 Error"])
set_row(t, 4, ["DMA 端到端强制性能目标", "通过", "20.62× / 33.84× / 21.11× / 37.28×"])
for row in (2, 4):
    t.rows[row].cells[1].paragraphs[0].runs[0].font.color.rgb = GREEN
    t.rows[row].cells[1].paragraphs[0].runs[0].bold = True

# Implementation table.
t = doc.tables[2]
set_row(t, 1, ["RTL 回归", "23/23 testbench + 协议检查 PASS", "通过"])
set_row(t, 5, ["实现时序", "100 MHz WNS/WHS +0.012/+0.010 ns；150 MHz +0.460/+0.060 ns", "通过"])
set_row(t, 6, ["整机资源", "31133 LUT、16707 FF、21 BRAM、78 DSP", "通过"])
set_row(t, 9, ["DRC/布线", "0 Error；0 未布线网络；Methodology Error=0", "通过"])
set_row(t, 10, ["最终板级自检", "DMA 与旧 AXI-Lite 综合自检均全部 PASS", "通过"])
for row in (1, 5, 6, 9, 10):
    t.rows[row].cells[2].paragraphs[0].runs[0].font.color.rgb = GREEN
    t.rows[row].cells[2].paragraphs[0].runs[0].bold = True

# Replace the obsolete “DMA still required” paragraph.
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("计入单样本 AXI-Lite 装载后"):
        replace_paragraph(
            paragraph,
            "原 AXI-Lite 单样本传输瓶颈已通过 DMA、突发传输、批处理和片上数据复用解决。"
            "最终板测单次共享查询 Tanimoto N64 为 20.62×（32 µs），8×N64 吞吐为 26.13×；"
            "GNN 为 33.84×，ADMET N64 为 21.11×，Pipeline 为 37.28×；"
            "Tanimoto 纯核心为 5 周期、120.00×。",
        )
    elif paragraph.text.startswith("完整日志："):
        replace_paragraph(
            paragraph,
            r"兼容性日志：D:\FPGA\reports\dma_batch\legacy-comprehensive-final.txt",
        )

# Mark the old P2 optimization item complete.
set_row(
    doc.tables[5], 6,
    ["已完成", "DMA、突发传输、批处理与片上复用", "最终板测和性能门禁全部通过"],
)

# Final deliverables table.
t = doc.tables[6]
delivery_rows = [
    ["类型", "路径"],
    ["正式比特流", r"D:\FPGA\artifacts\system_wrapper_dma_batch.bit"],
    ["正式硬件平台", r"D:\FPGA\artifacts\system_wrapper_dma_batch.xsa"],
    ["DMA 板级日志", r"D:\FPGA\reports\dma_batch\board-results.txt"],
    ["旧接口兼容日志", r"D:\FPGA\reports\dma_batch\legacy-comprehensive-final.txt"],
    ["实现门禁", r"D:\FPGA\reports\dma_batch\impl\gate_metrics.txt"],
    ["最终性能报告", r"D:\FPGA\reports\dma_batch\performance.md"],
    ["发布哈希", r"D:\FPGA\reports\dma_batch\release-hashes.sha256"],
    ["Release ELF", r"D:\FPGA\vitis_workspace\accelerator_dma_batch\Release\accelerator_dma_batch.elf"],
]
for index, values in enumerate(delivery_rows):
    set_row(t, index, values)

# Append the DMA acceptance supplement using the existing heading system.
doc.add_page_break()
doc.add_heading("8. DMA、突发传输与批处理架构", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "PS DDR → AXI DMA（Simple Mode）→ 128-bit AXIS FIFO → "
    "64 项混合任务队列 → Tanimoto / GNN / ADMET / Pipeline"
)
r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
r.font.size = Pt(10.5)

doc.add_paragraph(
    "控制面保持原 GP0/AXI-Lite 地址 0x43C00000；DMA 控制位于 0x80400000。"
    "MM2S/S2MM 存储侧为 128-bit，经 128→64 位宽转换器、AXI Interconnect 接入 PS HP0，"
    "数据域运行于 150 MHz；加速器与流接口运行于 100 MHz。输入采用 MOLQ 批头和八字任务头，"
    "输出采用 MOLR 记录和 MOLE 尾部，支持最多 64 个混合任务、2 MiB 单次传输、继续/停止错误策略、"
    "超时和确定性恢复。"
)
doc.add_paragraph(
    "为消除共享查询 Tanimoto 的短任务开销，软件使用独立 1 MiB 非缓存 TX 区；RTL 前端在消费"
    "当前 128-bit beat 第四个 32-bit 字的同一周期接收下一 beat，使稳态输入达到每四周期一拍。"
    "后端在核心计算前一候选时并行装载下一候选，隐藏候选间核心等待周期。"
)

doc.add_heading("9. 最终 DMA 性能与可靠性", level=1)
content_width = section_content_width_dxa(doc.sections[-1])
perf_rows = [
    ["项目", "实测", "目标", "结果"],
    ["MM2S（2 MiB）", "1198.37 MB/s", "≥500 MB/s", "PASS"],
    ["S2MM（2 MiB）", "1199.74 MB/s", "≥500 MB/s", "PASS"],
    ["单次共享 Tanimoto N64", "32 µs；20.62×", "≥20×", "PASS"],
    ["共享 Tanimoto 8×N64", "202 µs；26.13×", "吞吐证据", "PASS"],
    ["Tanimoto 纯核心", "5 周期；0.05 µs；120×", ">50×", "PASS"],
    ["GNN summary/full", "796 µs；33.84×", ">10×", "PASS"],
    ["ADMET N64", "97 µs；21.11×", ">20×", "PASS"],
    ["Pipeline 三模式", "1087 µs；37.28×", ">30×", "PASS"],
]
table = doc.add_table(rows=len(perf_rows), cols=4)
style_new_table(
    table,
    column_widths_from_weights([3.4, 2.2, 1.5, 1.0], content_width),
    perf_rows,
)

doc.add_paragraph(
    "单次 N64 相位剖析：TX/RX flush 0 µs、DMA+队列+核心 29 µs、RX invalidate 0 µs、"
    "响应解析 1 µs，总计 32 µs；该独立请求直接用于 ≥20× 强制验收。8×N64 总计 202 µs，"
    "只作为持续吞吐证据。任务记录的 post-payload 周期仅为队列尾延迟；纯核心 5 周期来自"
    " tb_tanimoto_latency 的直接测量。"
)

doc.add_page_break()
validation_rows = [
    ["验收项", "状态", "证据摘要"],
    ["功能与混合任务", "PASS", "四类任务、N64 与 0/1/2/3 顺序全部通过"],
    ["协议与异常恢复", "PASS", "continue/stop、timeout、DMA reset recovery"],
    ["压力测试", "PASS", "连续 1000 批；哈希 0x96FF4BF5"],
    ["实现门禁", "PASS", "双时钟域正时序；0 DRC/方法学错误"],
    ["旧接口兼容", "PASS", "ALL COMPREHENSIVE SELF-TESTS PASSED"],
    ["自动回归", "PASS", "23 RTL + 1 协议 + 7 主机端测试"],
]
table = doc.add_table(rows=len(validation_rows), cols=3)
style_new_table(
    table,
    column_widths_from_weights([2.3, 1.0, 4.4], content_width),
    validation_rows,
)

doc.add_heading("10. 正式发布与范围说明", level=1)
doc.add_paragraph(
    "正式 bit SHA-256：81EF25EEAA2ADB1783A28EAADCF305367E00756319CCEB973B5AD308FF9ABC3E。"
)
doc.add_paragraph(
    "正式 XSA SHA-256：736397FA243D1A30CBDAD151C1C3CFA19A0B0D1168BA68FE027A55ABFDEB6A61。"
)
doc.add_paragraph(
    "候选与正式产物哈希一致，已验证的旧 ILA 产物未被覆盖。至此十二项 DMA/突发传输/批处理任务全部完成。"
    "后续真实模型误差、MolRecommender 多智能体业务闭环、TCP 服务和 10 万/100 万分子压力测试仍需"
    "相应模型、数据与后端接口后另行验收。"
)

doc.core_properties.title = "Z15 FPGA 分子计算加速器验收与测试报告（DMA 最终版）"
doc.core_properties.subject = "DMA、突发传输、批处理、板级性能与发布验收"
doc.core_properties.comments = "Final DMA acceptance update, 2026-08-11"
doc.save(DST)
print(DST)
